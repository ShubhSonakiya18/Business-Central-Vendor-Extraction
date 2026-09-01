"""
V2 Document Loader
==================
Turns any supported input file into a `Document` of `TextSpan`s.

Dispatch:
    .pdf              -> per page: native text layer if present, else OCR
    .png/.jpg/.tif/... -> OCR
    .docx             -> python-docx paragraphs + tables, plus OCR of any
                         embedded images

Why text-layer-first for PDFs
-----------------------------
GST REG-06 and many portal-issued PDFs carry a real text layer: reading it is
exact, ~1000x faster than OCR, and immune to OCR confusions (O/0, I/1, S/5)
that matter enormously for GSTIN/PAN/IFSC. Scanned documents (cheques, some
Udyam certificates) have no text layer, so those pages fall back to OCR. The
decision is per page, because a single PDF can mix both. Pass
`force_ocr=True` to bypass the text layer entirely (useful for A/B comparison).

Every path emits the same TextSpan shape in the same pixel space, so the
layout and semantic engines never branch on input format.
"""

from __future__ import annotations

import io
import logging
import time
import zipfile
from pathlib import Path
from typing import Iterable, Optional

import numpy as np

from ..models import (
    PDF_POINTS_PER_INCH,
    RENDER_DPI,
    BBox,
    Document,
    DocumentSet,
    DocumentType,
    Page,
    SpanSource,
    TableRef,
    TextSpan,
)
from .ocr_engine import OCREngine, sort_reading_order

logger = logging.getLogger(__name__)


def _default_dpi_for(engine: Optional[OCREngine]) -> int:
    """The right render DPI for whichever engine will actually run.

    RENDER_DPI (200) and RAPID_RENDER_DPI (100, paired with
    intra_op_num_threads=8 -- see settings.py) were each tuned for their own
    engine; using the wrong one is a silent latency/accuracy mismatch, not an
    error, so callers must not have to know which engine is active to get
    this right. `OCREngine()` construction is cheap (no model load), so
    probing `.backend` here when no engine was supplied costs nothing.
    """
    from ....config.settings import RAPID_RENDER_DPI

    backend = (engine or OCREngine()).backend
    return RAPID_RENDER_DPI if backend == "rapidocr" else RENDER_DPI

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}

# A page whose text layer yields fewer than this many characters is treated as
# scanned. Digital pages carry hundreds; image-only pages carry 0 or a few
# stray characters from a stamp/annotation.
MIN_TEXT_LAYER_CHARS = 40

# Fraction of the page covered by raster images above which the page is
# considered a scan, and its text layer is distrusted even when populated.
#
# This exists because character count alone is not enough. The sample cancelled
# cheque is a 100%-image page carrying a 536-character text layer produced by
# some upstream tool, and that layer is mangled -- "ICICI Bank" comes through
# as "Otctctean", "Hindustan Road" as "Hlndu.tan Rord". Trusting it silently
# corrupted exactly the fields that matter most (bank name, branch, IFSC).
# Re-OCRing the render recovers them correctly.
#
# Real digital pages are far below this line (the GST certificate sits at
# ~1%, its images being only a logo and a signature), so the margin is wide.
MAX_IMAGE_COVERAGE_FOR_TEXT_LAYER = 0.5


# ---------------------------------------------------------------------------
# PUBLIC API
# ---------------------------------------------------------------------------

def load_document(
    path: str | Path,
    engine: Optional[OCREngine] = None,
    force_ocr: bool = False,
    dpi: Optional[int] = None,
) -> Document:
    """Read one file into the common representation.

    `dpi=None` (the default) resolves the same way as `load_documents()` --
    see `_default_dpi_for()` -- so calling this directly does not silently
    render at the wrong engine's DPI.
    """
    if dpi is None:
        dpi = _default_dpi_for(engine)
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path, engine=engine, force_ocr=force_ocr, dpi=dpi)
    if suffix in IMAGE_SUFFIXES:
        return _load_image(path, engine=engine)
    if suffix == ".docx":
        return _load_docx(path, engine=engine)
    raise ValueError(f"Unsupported file type: {path.name} ({suffix})")


def load_documents(
    paths: Iterable[str | Path],
    engine: Optional[OCREngine] = None,
    force_ocr: bool = False,
    dpi: Optional[int] = None,
    skip_unsupported: bool = True,
) -> DocumentSet:
    """Read many files. One vendor generally means several documents, and the
    semantic engine merges candidates across all of them.

    `dpi=None` (the default) resolves to whichever DPI matches the active
    engine's backend -- see `_default_dpi_for()`. Pass an explicit `dpi=` to
    override that for a benchmark/A-B run; it is used as given, unchanged.
    """
    if dpi is None:
        dpi = _default_dpi_for(engine)

    docs: list[Document] = []
    for p in paths:
        p = Path(p)
        try:
            docs.append(load_document(p, engine=engine, force_ocr=force_ocr, dpi=dpi))
        except ValueError as exc:
            if not skip_unsupported:
                raise
            logger.warning("Skipping %s: %s", p.name, exc)

    # Phase 0 instrumentation (plan.md 0.2): how often the native-text fast
    # path fires vs falling through to OCR. Logged, not persisted -- this is
    # a diagnostic for the benchmark protocol, not part of the data model.
    pages = [pg for d in docs for pg in d.pages]
    if pages:
        text_hits = sum(1 for pg in pages if pg.extraction_method == "pdf_text")
        logger.info(
            "_should_ocr fast path: %d/%d pages used the text layer (%.0f%%); "
            "%d fell through to OCR",
            text_hits, len(pages), 100 * text_hits / len(pages), len(pages) - text_hits,
        )

    return DocumentSet(documents=docs)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def _load_pdf(path: Path, engine: Optional[OCREngine], force_ocr: bool, dpi: int) -> Document:
    import pypdfium2 as pdfium

    scale = dpi / PDF_POINTS_PER_INCH
    doc = Document(path=str(path), doc_type=DocumentType.PDF)

    # Phase 0 instrumentation (plan.md 0.1): pdf_open_s and render_s are stored
    # in doc.metadata, a pre-existing free-form dict, so this is purely
    # additive -- no change to Page/Document's shape or serialization.
    open_t0 = time.perf_counter()
    pdf = pdfium.PdfDocument(str(path))
    pdf_open_s = time.perf_counter() - open_t0
    render_s = 0.0

    try:
        for index, page in enumerate(pdf):
            t0 = time.perf_counter()
            page_no = index + 1
            width_pt, height_pt = page.get_size()
            page_obj = Page(
                number=page_no,
                width=width_pt * scale,
                height=height_pt * scale,
                extraction_method="ocr" if force_ocr else "pdf_text",
            )

            spans: list[TextSpan] = []
            reason = "forced"
            if not force_ocr:
                spans = _pdf_text_spans(
                    page=page,
                    page_no=page_no,
                    page_height_pt=height_pt,
                    scale=scale,
                    source_document=path.name,
                )
                use_ocr, reason = _should_ocr(page, spans)
            else:
                use_ocr = True

            if use_ocr:
                if engine is None:
                    engine = OCREngine()
                render_t0 = time.perf_counter()
                image = _render_page(page, scale)
                render_s += time.perf_counter() - render_t0
                spans = engine.read_image(
                    image, source_document=path.name, page=page_no, span_source=SpanSource.OCR
                )
                page_obj.extraction_method = "ocr"
                page_obj.ocr_reason = reason
            else:
                spans = sort_reading_order(spans)
                for i, s in enumerate(spans):
                    s.order = i
                page_obj.extraction_method = "pdf_text"

            page_obj.spans = spans
            page_obj.duration_s = time.perf_counter() - t0
            doc.pages.append(page_obj)
    finally:
        pdf.close()

    methods = {p.extraction_method for p in doc.pages}
    doc.metadata["extraction_methods"] = sorted(methods)
    doc.metadata["dpi"] = dpi
    doc.metadata["pdf_open_s"] = round(pdf_open_s, 4)
    doc.metadata["render_s"] = round(render_s, 4)
    return doc


def _render_page(page, scale: float) -> np.ndarray:
    bitmap = page.render(scale=scale)
    return bitmap.to_numpy()


def _should_ocr(page, spans: list[TextSpan]) -> tuple[bool, str]:
    """Decide whether this page's text layer can be trusted.

    Returns (use_ocr, reason). The reason is recorded on the Page so a later
    audit can explain why a given document was read the way it was.
    """
    if len(_joined(spans)) < MIN_TEXT_LAYER_CHARS:
        return True, "no_text_layer"

    coverage = _image_coverage(page)
    if coverage >= MAX_IMAGE_COVERAGE_FOR_TEXT_LAYER:
        return True, f"scanned_page(image_coverage={coverage:.0%})"

    return False, ""


def _image_coverage(page) -> float:
    """Fraction of the page area covered by raster image objects.

    A digital PDF uses images only for logos and signatures (a few percent);
    a scan is one image covering the whole page.
    """
    import pypdfium2.raw as praw

    try:
        width, height = page.get_size()
        page_area = float(width) * float(height)
        if page_area <= 0:
            return 0.0

        covered = 0.0
        for obj in page.get_objects():
            if obj.type != praw.FPDF_PAGEOBJ_IMAGE:
                continue
            try:
                left, bottom, right, top = obj.get_bounds()
            except Exception:
                continue
            covered += abs((right - left) * (top - bottom))
        # Overlapping images can sum past the page area; clamp rather than
        # reporting a nonsensical >100% coverage.
        return min(covered / page_area, 1.0)
    except Exception as exc:
        logger.warning("Could not measure image coverage: %s", exc)
        return 0.0


def _pdf_text_spans(
    page,
    page_no: int,
    page_height_pt: float,
    scale: float,
    source_document: str,
) -> list[TextSpan]:
    """Extract native text as line-level spans with real geometry.

    Works character-by-character rather than using pdfium's rect API: pdfium
    sometimes merges visually distant text into one rect, which would destroy
    the label/value adjacency the semantic engine depends on. Grouping chars
    ourselves also lets us split a visual line wherever there is a wide gap,
    mirroring how OCR emits label and value as separate boxes.
    """
    textpage = page.get_textpage()
    try:
        n_chars = textpage.count_chars()
        if n_chars == 0:
            return []

        chars: list[tuple[str, BBox]] = []
        for i in range(n_chars):
            try:
                ch = textpage.get_text_range(i, 1)
            except Exception:
                continue
            if not ch or ch in "\r\n":
                chars.append((ch or "\n", None))  # newline marker, no geometry
                continue
            try:
                left, bottom, right, top = textpage.get_charbox(i)
            except Exception:
                continue
            # PDF space is bottom-left origin; flip to top-left and scale to px.
            bbox = BBox(
                x1=left * scale,
                y1=(page_height_pt - top) * scale,
                x2=right * scale,
                y2=(page_height_pt - bottom) * scale,
            )
            chars.append((ch, bbox))
    finally:
        textpage.close()

    return _group_chars_into_spans(chars, page_no, source_document)


def _group_chars_into_spans(
    chars: list[tuple[str, Optional[BBox]]],
    page_no: int,
    source_document: str,
) -> list[TextSpan]:
    """Group characters into line segments, splitting on newlines, vertical
    steps, and wide horizontal gaps."""
    widths = [b.width for _, b in chars if b is not None and b.width > 0]
    median_w = float(np.median(widths)) if widths else 4.0
    gap_threshold = median_w * 2.5

    spans: list[TextSpan] = []
    buf_text: list[str] = []
    buf_boxes: list[BBox] = []
    prev: Optional[BBox] = None

    def flush() -> None:
        nonlocal buf_text, buf_boxes
        text = "".join(buf_text).strip()
        if text and buf_boxes:
            spans.append(
                TextSpan(
                    text=text,
                    page=page_no,
                    bbox=BBox(
                        min(b.x1 for b in buf_boxes),
                        min(b.y1 for b in buf_boxes),
                        max(b.x2 for b in buf_boxes),
                        max(b.y2 for b in buf_boxes),
                    ),
                    source_document=source_document,
                    confidence=1.0,  # native text is exact
                    source=SpanSource.PDF_TEXT,
                )
            )
        buf_text, buf_boxes = [], []

    for ch, bbox in chars:
        if bbox is None:  # newline
            flush()
            prev = None
            continue

        if prev is not None:
            same_line = prev.vertical_overlap(bbox) >= 0.5
            gap = bbox.x1 - prev.x2
            if not same_line or gap > gap_threshold:
                flush()
        buf_text.append(ch)
        buf_boxes.append(bbox)
        prev = bbox

    flush()
    return spans


# ---------------------------------------------------------------------------
# IMAGE
# ---------------------------------------------------------------------------

def _load_image(path: Path, engine: Optional[OCREngine]) -> Document:
    from PIL import Image

    if engine is None:
        engine = OCREngine()

    t0 = time.perf_counter()
    with Image.open(path) as im:
        im = im.convert("RGB")
        array = np.asarray(im)

    spans = engine.read_image(array, source_document=path.name, page=1)
    page = Page(
        number=1,
        width=float(array.shape[1]),
        height=float(array.shape[0]),
        spans=spans,
        extraction_method="ocr",
        duration_s=time.perf_counter() - t0,
    )
    doc = Document(path=str(path), doc_type=DocumentType.IMAGE, pages=[page])
    doc.metadata["extraction_methods"] = ["ocr"]
    return doc


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

# DOCX carries no usable page geometry, so we synthesize a layout on an
# A4-sized canvas at RENDER_DPI. The absolute numbers are meaningless; what
# matters is that they preserve the *relations* the semantic engine reads --
# a table value sits to the right of its label, a paragraph below the one
# before it. Spans produced this way are flagged synthetic_bbox=True.
DOCX_PAGE_W = 8.27 * RENDER_DPI
DOCX_PAGE_H = 11.69 * RENDER_DPI
DOCX_MARGIN = 0.75 * RENDER_DPI
DOCX_LINE_H = 0.22 * RENDER_DPI
DOCX_CHAR_W = 0.075 * RENDER_DPI


def _load_docx(path: Path, engine: Optional[OCREngine]) -> Document:
    import docx

    t0 = time.perf_counter()
    d = docx.Document(str(path))
    doc = Document(path=str(path), doc_type=DocumentType.DOCX)

    cursor = _DocxCursor(source_document=path.name)
    for block in _iter_docx_blocks(d):
        if block[0] == "p":
            cursor.add_paragraph(block[1])
        else:
            cursor.add_table(block[1], block[2])

    pages = cursor.finish()

    # Embedded images (a scanned certificate pasted into a DOCX) still need
    # OCR -- but only the images, not a rasterization of the whole document.
    image_spans_by_page: list[TextSpan] = []
    images = _docx_images(path)
    if images:
        if engine is None:
            engine = OCREngine()
        for idx, blob in enumerate(images):
            array = _blob_to_array(blob)
            if array is None:
                continue
            spans = engine.read_image(
                array,
                source_document=path.name,
                page=len(pages) + 1,
                span_source=SpanSource.DOCX_IMAGE_OCR,
            )
            for s in spans:
                s.bbox = BBox(s.bbox.x1, s.bbox.y1, s.bbox.x2, s.bbox.y2)
            if spans:
                image_spans_by_page.append(
                    Page(
                        number=len(pages) + 1,
                        width=float(array.shape[1]),
                        height=float(array.shape[0]),
                        spans=spans,
                        extraction_method="docx_image_ocr",
                    )
                )
                pages.append(image_spans_by_page[-1])
                logger.info("OCR'd embedded image %d in %s: %d spans", idx, path.name, len(spans))

    doc.pages = pages
    if doc.pages:
        doc.pages[0].duration_s = time.perf_counter() - t0
    doc.metadata["extraction_methods"] = sorted({p.extraction_method for p in doc.pages})
    doc.metadata["embedded_images"] = len(images)
    return doc


def _iter_docx_blocks(document):
    """Yield paragraphs and tables in true document order.

    python-docx exposes `.paragraphs` and `.tables` as separate flat lists,
    which loses their interleaving -- and in a vendor form the order is
    exactly what tells you which heading a table belongs to. So we walk the
    body XML instead.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield ("p", Paragraph(child, document))
        elif tag == "tbl":
            table = Table(child, document)
            yield ("t", table, None)


class _DocxCursor:
    """Lays synthesized text out top-to-bottom, paginating on overflow."""

    def __init__(self, source_document: str):
        self.source_document = source_document
        self.pages: list[Page] = []
        self._new_page()

    def _new_page(self) -> None:
        self.page = Page(
            number=len(self.pages) + 1,
            width=DOCX_PAGE_W,
            height=DOCX_PAGE_H,
            extraction_method="docx",
        )
        self.pages.append(self.page)
        self.y = DOCX_MARGIN
        self.order = 0

    def _advance(self, height: float) -> None:
        self.y += height
        if self.y > DOCX_PAGE_H - DOCX_MARGIN:
            self._new_page()

    def _emit(self, text: str, x1: float, x2: float, source: SpanSource, table: Optional[TableRef]) -> None:
        span = TextSpan(
            text=text,
            page=self.page.number,
            bbox=BBox(x1, self.y, x2, self.y + DOCX_LINE_H),
            source_document=self.source_document,
            confidence=1.0,
            source=source,
            order=self.order,
            table=table,
            synthetic_bbox=True,
        )
        self.page.spans.append(span)
        self.order += 1

    def add_paragraph(self, paragraph) -> None:
        text = (paragraph.text or "").strip()
        if not text:
            return
        width = min(len(text) * DOCX_CHAR_W, DOCX_PAGE_W - 2 * DOCX_MARGIN)
        self._emit(text, DOCX_MARGIN, DOCX_MARGIN + width, SpanSource.DOCX_PARAGRAPH, None)
        self._advance(DOCX_LINE_H)

    def add_table(self, table, table_index: Optional[int]) -> None:
        idx = table_index if table_index is not None else 0
        rows = table.rows
        if not rows:
            return
        n_cols = max(len(r.cells) for r in rows)
        col_w = (DOCX_PAGE_W - 2 * DOCX_MARGIN) / max(n_cols, 1)

        for r_i, row in enumerate(rows):
            row_top = self.y
            for c_i, cell in enumerate(row.cells):
                text = (cell.text or "").strip()
                if not text:
                    continue
                x1 = DOCX_MARGIN + c_i * col_w
                self.y = row_top
                self._emit(
                    text,
                    x1,
                    x1 + col_w,
                    SpanSource.DOCX_TABLE,
                    TableRef(table_index=idx, row=r_i, col=c_i),
                )
            self.y = row_top
            self._advance(DOCX_LINE_H)

    def finish(self) -> list[Page]:
        return [p for p in self.pages if p.spans]


def _docx_images(path: Path) -> list[bytes]:
    """Pull embedded image blobs straight from the DOCX zip container."""
    blobs: list[bytes] = []
    try:
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/") and Path(name).suffix.lower() in IMAGE_SUFFIXES:
                    blobs.append(zf.read(name))
    except zipfile.BadZipFile:
        logger.warning("%s is not a readable DOCX container", path.name)
    return blobs


def _blob_to_array(blob: bytes) -> Optional[np.ndarray]:
    from PIL import Image

    try:
        with Image.open(io.BytesIO(blob)) as im:
            # Tiny images are logos/bullets, not scanned content.
            if im.width < 80 or im.height < 80:
                return None
            return np.asarray(im.convert("RGB"))
    except Exception as exc:
        logger.warning("Unreadable embedded image: %s", exc)
        return None


def _joined(spans: list[TextSpan]) -> str:
    return "".join(s.text for s in spans).strip()
